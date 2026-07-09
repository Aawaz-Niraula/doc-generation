"""render_docx.py — flowing Word document renderer (python-docx).

Consumes the same DocumentSpec as render_pdf / render_pptx. Interactivity:
live Table of Contents field + paragraph outline levels (Navigation pane),
running footer with a live PAGE field, auto field refresh on open.

Design-system rules enforced here: metric-safe Office fonts from the theme,
heading bands as tinted cells (background tint, not decorative bars — the
cover's old accent bar is gone), pull quotes as tinted blocks instead of
single-edge borders.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from content_model import DocumentSpec, Section, TableBlock
from theme import Theme, hex_to_rgb, mix, readable_on


def _hexrgb(h: str) -> RGBColor:
    return RGBColor(*hex_to_rgb(h))


def _run(run, font=None, size=None, color=None, bold=None, italic=None,
         caps=False, spacing=None):
    if font:
        run.font.name = font
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rfonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _hexrgb(color) if isinstance(color, str) else color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if caps:
        run.font.all_caps = True
    if spacing is not None:  # letter-spacing in twips
        rpr = run._r.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing)); rpr.append(sp)
    return run


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _strip_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'none'); borders.append(b)
    tblPr.append(borders)
    return tblPr


def _set_outline_level(paragraph, level: int):
    """Tag a paragraph with a Word outline level so it appears in the Navigation
    pane and is picked up by the TOC field (\\u switch)."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level))
    pPr.append(ol)


def _add_toc(doc: Document):
    """Live, clickable Table of Contents field (refreshes on open)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    r = run._r
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Table of contents — right-click and choose “Update Field”."
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for child in (begin, instr, sep, placeholder, end):
        r.append(child)


def _enable_update_fields(doc: Document):
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        el = OxmlElement('w:updateFields')
        el.set(qn('w:val'), 'true')
        settings.append(el)


def _heading_band(doc: Document, text: str, bg_hex: str, size_pt: float = 14,
                  outline_level: Optional[int] = None, font: str = "Calibri"):
    """Section heading on a tinted band (functional background tint)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    _set_cell_bg(cell, bg_hex)
    _strip_table_borders(table)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Inches(0.16)
    run = p.add_run(text)
    _run(run, font=font, size=size_pt, bold=True, color=readable_on(bg_hex))
    if outline_level is not None:
        _set_outline_level(p, outline_level)
    doc.add_paragraph()


def _pull_quote(doc: Document, text: str, theme: Theme):
    """Pull quote as a tinted rounded-feel block (no single-edge border)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    tint = mix(theme.dominant, "#FFFFFF", 0.92)
    _set_cell_bg(cell, tint)
    _strip_table_borders(table)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.22)
    run = p.add_run(f'“{text}”')
    _run(run, font=theme.pptx_fonts["display"], size=12.5, italic=True,
         color=theme.dominant)
    doc.add_paragraph()


def _data_table(doc: Document, table: TableBlock, theme: Theme):
    """Banded data table with alternating row tints."""
    headers = [str(h) for h in table.headers[:6]]
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.style = 'Table Grid'
    t.autofit = True
    hdr_cells = t.rows[0].cells
    for ci, h in enumerate(headers):
        _set_cell_bg(hdr_cells[ci], theme.dominant)
        para = hdr_cells[ci].paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(h)
        _run(run, font=theme.pptx_fonts["data"], size=10, bold=True,
             color=readable_on(theme.dominant))
    alt = mix(theme.dominant, "#FFFFFF", 0.90)
    for ri, row in enumerate(table.rows[:12]):
        cells = t.add_row().cells
        vals = (list(row) + [""] * n)[:n]
        for ci, v in enumerate(vals):
            if ri % 2 == 1:
                _set_cell_bg(cells[ci], alt)
            para = cells[ci].paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(str(v))
            _run(run, size=10, bold=(ci == 0),
                 color=theme.dominant if ci == 0 else "#1F2A3C")
    doc.add_paragraph()


def _stats_table(doc: Document, section: Section, theme: Theme):
    """Render stat cards as a compact metric table when a section carries them."""
    stats = [s for s in section.stats if s.value][:4]
    if not stats:
        return
    block = TableBlock(headers=["Metric", "Value", "Context"],
                       rows=[[s.label or f"Metric {i+1}", s.value, s.context]
                             for i, s in enumerate(stats)])
    _data_table(doc, block, theme)


def _set_base_fonts(doc: Document, body_font: str, body_pt: float = 10.5,
                    body_color: str = "26303F"):
    """Default (Normal) typeface so every paragraph inherits the theme body font."""
    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(body_pt)
    style.font.color.rgb = _hexrgb(body_color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), body_font)
    style.paragraph_format.line_spacing = 1.35
    style.paragraph_format.space_after = Pt(6)


def _page_footer(doc: Document, left_text: str, theme: Theme):
    """Running footer: document title left, live page number right."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab'); tab.set(qn('w:val'), 'right')
    content_w = int((section.page_width - section.left_margin - section.right_margin) / 635)
    tab.set(qn('w:pos'), str(content_w)); tabs.append(tab); pPr.append(tabs)
    data_font = theme.pptx_fonts["data"]
    accent = theme.accent_readable_on_light()
    lr = p.add_run(left_text.upper())
    _run(lr, font=data_font, size=7.5, color="9AA2AD", spacing=20)
    p.add_run("\t")
    pr = p.add_run("PAGE ")
    _run(pr, font=data_font, size=7.5, color=accent)
    fld = p.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    fld._r.append(begin); fld._r.append(instr); fld._r.append(end)
    _run(fld, font=data_font, size=7.5, color=accent, bold=True)


# ─── Entry point ──────────────────────────────────────────────────────────────

def render_docx(spec: DocumentSpec, theme: Theme) -> bytes:
    doc = Document()
    F = theme.pptx_fonts       # metric-safe Office fonts for Word too
    _set_base_fonts(doc, F["body"], body_pt=10.5)

    for sec in doc.sections:
        sec.top_margin = Inches(1.15)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.15)
        sec.right_margin = Inches(1.15)

    body_sections = [s for s in spec.sections if s.kind not in ("cover", "closing")]
    cover = next((s for s in spec.sections if s.kind == "cover"), None)
    closing = next((s for s in spec.sections if s.kind == "closing"), None)
    accent_light = theme.accent_readable_on_light()

    # ── editorial cover (hierarchy by whitespace — no accent bar) ─────────────
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(90)

    eyebrow_para = doc.add_paragraph()
    eyebrow_para.paragraph_format.space_after = Pt(10)
    _run(eyebrow_para.add_run((theme.aesthetic_label or "Report").upper()),
         font=F["data"], size=10.5, color=accent_light, bold=True, spacing=48)

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.line_spacing = 1.0
    _run(title_para.add_run(spec.title), font=F["display"], size=38,
         color=theme.dominant, bold=True)

    subtitle = spec.subtitle or (cover.subtitle if cover else "")
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(26)
        _run(sp.add_run(subtitle), font=F["heading"], size=15, color="55606E")

    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_before = Pt(20)
    meta_line = "  ·  ".join(t for t in (spec.author, spec.date) if t)
    _run(meta_para.add_run(meta_line), font=F["data"], size=9.5,
         color="6B7280", spacing=20)

    doc.add_page_break()

    # ── contents ──────────────────────────────────────────────────────────────
    if body_sections:
        _heading_band(doc, "Contents", theme.dominant, size_pt=13, font=F["heading"])
        _add_toc(doc)
        doc.add_page_break()

    # ── executive summary ─────────────────────────────────────────────────────
    abstract = spec.abstract or (cover.body if cover else "")
    if abstract:
        _heading_band(doc, "Executive Summary", theme.dominant, size_pt=13,
                      outline_level=0, font=F["heading"])
        abs_para = doc.add_paragraph(abstract)
        abs_para.paragraph_format.space_after = Pt(12)
        abs_para.paragraph_format.left_indent = Inches(0.28)
        abs_para.paragraph_format.line_spacing = 1.4
        for run in abs_para.runs:
            _run(run, font=F["heading"], size=12.5, color="3A4453")
        doc.add_paragraph()

    # ── body sections ─────────────────────────────────────────────────────────
    for si, section in enumerate(body_sections):
        band_bg = theme.secondary if si % 2 == 1 else theme.dominant
        _heading_band(doc, f"{si + 1:02d}   {section.heading}", band_bg,
                      size_pt=14, outline_level=0, font=F["heading"])

        if section.body:
            body_para = doc.add_paragraph(section.body)
            body_para.paragraph_format.space_after = Pt(11)
            body_para.paragraph_format.space_before = Pt(3)
            body_para.paragraph_format.line_spacing = 1.4
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in body_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = _hexrgb("#26303F")

        if section.callout:
            _pull_quote(doc, section.callout, theme)

        if section.table and not section.table.is_empty():
            _data_table(doc, section.table, theme)
        elif section.stats:
            _stats_table(doc, section, theme)

        if section.steps:
            for i, step in enumerate(section.steps[:6]):
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)
                sp.paragraph_format.left_indent = Inches(0.2)
                lead = " · ".join(t for t in (step.date, step.title) if t)
                _run(sp.add_run(f"{i + 1:02d}  {lead}   "), font=F["data"], size=10.5,
                     bold=True, color=accent_light)
                _run(sp.add_run(step.description), size=10.5, color="#26303F")
            doc.add_paragraph()

        for sub_idx, sub in enumerate(section.subsections):
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.space_before = Pt(11)
            sub_para.paragraph_format.space_after = Pt(4)
            _run(sub_para.add_run(f"{si + 1}.{sub_idx + 1}   {sub.heading}"),
                 font=F["heading"], size=13, bold=True, color=accent_light)
            _set_outline_level(sub_para, 1)
            sub_body = doc.add_paragraph(sub.body)
            sub_body.paragraph_format.space_after = Pt(8)
            sub_body.paragraph_format.left_indent = Inches(0.2)
            for run in sub_body.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = _hexrgb("#26303F")

        doc.add_paragraph()   # whitespace between sections (no separator rules)

    # ── conclusion ────────────────────────────────────────────────────────────
    conclusion = spec.conclusion or (closing.body if closing else "")
    if conclusion:
        _heading_band(doc, "Conclusion", theme.dominant, size_pt=13,
                      outline_level=0, font=F["heading"])
        conc_para = doc.add_paragraph(conclusion)
        conc_para.paragraph_format.space_after = Pt(10)
        conc_para.paragraph_format.line_spacing = 1.4
        for run in conc_para.runs:
            _run(run, font=F["heading"], size=11.5, color="3A4453")
        if closing and (closing.cta or closing.callout):
            _pull_quote(doc, closing.cta or closing.callout, theme)

    _page_footer(doc, spec.title, theme)
    _enable_update_fields(doc)

    core = doc.core_properties
    core.title = spec.title
    core.author = spec.author or "AI Document Studio"
    core.subject = spec.subtitle or spec.title
    core.comments = f"Theme: {theme.name} · {theme.aesthetic_label}"
    core.created = datetime.now(timezone.utc).replace(tzinfo=None)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
